import os
import uuid
import pandas as pd
from django.shortcuts import render
from django.conf import settings
from docx import Document
from .forms import ExcelUploadForm
from .services import AIService

def index(request):
    form = ExcelUploadForm()
    if request.method != 'POST':
        return render(request, 'excel_processor/index.html', {'form': form})

    form = ExcelUploadForm(request.POST, request.FILES)
    if not form.is_valid():
        return render(request, 'excel_processor/index.html', {'form': form})

    try:
        excel_file = request.FILES['excel_file']
        df = pd.read_excel(excel_file)
        
        df.columns = df.columns.str.strip()

        col_length = 'Протяженность, км'
        if col_length not in df.columns:
            raise ValueError(f"Колонка '{col_length}' не найдена в файле.")

        df[col_length] = pd.to_numeric(
            df[col_length].astype(str).str.replace(',', '.'), 
            errors='coerce'
        ).fillna(0)

        total_sum = round(df[col_length].sum(), 2)
        ai_result = AIService.get_llm_analysis(df, total_sum)

        # Создание документа
        doc = Document()
        doc.add_heading('Отчет по автомобильным дорогам', level=1)
        
        if ai_result:
            doc.add_heading('Аналитическое резюме', level=2)
            p = doc.add_paragraph(ai_result)
            p.alignment = 3

        doc = Document()
        doc.add_heading('Отчет по автомобильным дорогам', level=1)
        
        if ai_result:
            doc.add_heading('Аналитическое резюме', level=2)
            p = doc.add_paragraph(ai_result)
            p.alignment = 3
            doc.add_paragraph("") 

        doc.add_paragraph("Таблица 1 - Перечень и характеристика автомобильных дорог")
        
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        for i, column in enumerate(df.columns):
            table.rows[0].cells[i].text = str(column)
        
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

        doc.add_paragraph(f"\nОбщая протяженность автомобильных дорог составляет {total_sum} км")

        file_id = uuid.uuid4().hex[:8]
        file_name = f"Report_{file_id}.docx"
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)

        if not os.path.exists(settings.MEDIA_ROOT):
            os.makedirs(settings.MEDIA_ROOT)
            
        doc.save(file_path)

        return render(request, 'excel_processor/result.html', {
            'report_url': settings.MEDIA_URL + file_name,
            'summary': f"Итоговая протяженность: {total_sum} км"
        })

    except Exception as e:
        return render(request, 'excel_processor/index.html', {
            'form': form, 
            'error': f"Ошибка при обработке: {str(e)}"
        })